import docx

from app.core.exceptions import ExtractionError


def extract_text_from_docx(file_path: str) -> str:
    try:
        document = docx.Document(file_path)
    except Exception as exc:
        raise ExtractionError(f"Failed to read DOCX file: {exc}") from exc

    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    paragraphs.append(cell.text.strip())

    text = "\n".join(paragraphs).strip()
    if not text:
        raise ExtractionError("DOCX file does not contain any readable text")
    return text
